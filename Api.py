import os
import hashlib
from uuid import uuid4
from datetime import datetime
from typing import Dict, Any, List, Optional, Set

from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
from dotenv import load_dotenv

# --- Your existing modules ---
from processing import PDFProcessor, VectorStore
from document_cache import DocumentCache

# --- LangChain / Agentic bits ---
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.documents import Document
from langchain_core.retrievers import BaseRetriever
from langchain_core.callbacks import CallbackManagerForRetrieverRun
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain.tools.retriever import create_retriever_tool
from langchain.agents import AgentExecutor, create_tool_calling_agent

load_dotenv()

# -----------------------------------------------------------------------------
# Flask app setup
# -----------------------------------------------------------------------------

app = Flask(__name__)
CORS(app)

app.config["UPLOAD_FOLDER"] = "./uploads"
app.config["MAX_CONTENT_LENGTH"] = 2000 * 1024 * 1024  # 2 GB
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

# -----------------------------------------------------------------------------
# Global state (in-memory)
# -----------------------------------------------------------------------------

# Single global vector DB (your VectorStore using FastEmbed + BGE-small-en-v1.5 + FAISS)
vector_store = VectorStore()

# PDF processor + page cache
pdf_processor = PDFProcessor()
document_cache = DocumentCache(ttl=3600)

# Simple in-memory KB + documents registry
knowledge_bases: Dict[str, Dict[str, Any]] = {}
documents: Dict[str, Dict[str, Any]] = {}

# Per-conversation chat histories for the agent
_session_histories: Dict[str, ChatMessageHistory] = {}

# Cached LLM instance
_llm: Optional[ChatGoogleGenerativeAI] = None


def _now_iso() -> str:
    """Return current UTC time in ISO-8601 format."""
    return datetime.utcnow().isoformat() + "Z"


def _ensure_default_kb() -> str:
    """Create a 'default' KB if it doesn't exist and return its ID."""
    kb_id = "default"
    if kb_id not in knowledge_bases:
        now = _now_iso()
        knowledge_bases[kb_id] = {
            "id": kb_id,
            "name": "Default",
            "description": "Default Knowledge Base",
            "visibility": "private",
            "created_at": now,
            "updated_at": now,
            "document_ids": [],
        }
    return kb_id


DEFAULT_KB_ID = _ensure_default_kb()

# -----------------------------------------------------------------------------
# LangChain primitives – retriever + agent with memory
# -----------------------------------------------------------------------------

class KBVectorRetriever(BaseRetriever):
    """
    LangChain retriever wrapper around your custom VectorStore.

    It:
    - Uses vector_store.search(query, k)
    - Optionally filters by kb_ids
    - Returns LangChain Document objects with your metadata attached
    """

    # These are Pydantic model fields now
    vector_store: VectorStore
    kb_ids: Optional[Set[str]] = None
    k: int = 5

    class Config:
        # Allow VectorStore (a non-pydantic type) as a field
        arbitrary_types_allowed = True

    def _get_relevant_documents(
        self,
        query: str,
        *,
        run_manager: Optional[CallbackManagerForRetrieverRun] = None,
    ) -> List[Document]:
        # Convert to a real set once, for quick membership checks
        kb_id_set = self.kb_ids if self.kb_ids else None

        results = self.vector_store.search(query, k=self.k)
        docs: List[Document] = []

        for meta, dist in results:
            kb_id = meta.get("kb_id")

            # Optional KB filtering
            if kb_id_set and kb_id not in kb_id_set:
                continue

            rich_meta = dict(meta)
            rich_meta["score"] = dist

            docs.append(
                Document(
                    page_content=meta.get("doc_text", ""),
                    metadata=rich_meta,
                )
            )
        return docs

def _get_gemini_api_key() -> str:
    """
    Priority:
    1) GEMINI_API_KEY_AGENT
    2) GEMINI_API_KEY_QUERY
    3) GEMINI_API_KEY_GENERATE
    """
    api_key = (
        os.environ.get("GEMINI_API_KEY_AGENT")
        or os.environ.get("GEMINI_API_KEY_QUERY")
        or os.environ.get("GEMINI_API_KEY_GENERATE")
    )
    if not api_key:
        raise RuntimeError(
            "No Gemini API key found. "
            "Set GEMINI_API_KEY_AGENT or GEMINI_API_KEY_QUERY or GEMINI_API_KEY_GENERATE."
        )
    return api_key


def _get_llm() -> ChatGoogleGenerativeAI:
    """Singleton ChatGoogleGenerativeAI, used by the agent."""
    global _llm
    if _llm is None:
        api_key = _get_gemini_api_key()
        # langchain-google-genai reads GOOGLE_API_KEY
        os.environ.setdefault("GOOGLE_API_KEY", api_key)

        _llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            temperature=0.2,
            max_output_tokens=2048,
        )
    return _llm


def _get_session_history(session_id: str) -> ChatMessageHistory:
    """Return (and lazily create) a ChatMessageHistory for a given conversation."""
    if session_id not in _session_histories:
        _session_histories[session_id] = ChatMessageHistory()
    return _session_histories[session_id]


def _build_kb_agent_with_history(
    kb_ids: List[str],
    top_k: int = 5,
) -> RunnableWithMessageHistory:
    """
    Build an agent that:
    - Uses a retriever tool over your vector store
    - Has chat history per conversation_id
    - Returns answers grounded to retrieved docs
    """
    # 1. Retriever & tool
    retriever = KBVectorRetriever(vector_store=vector_store, kb_ids=kb_ids, k=top_k)
    retriever_tool = create_retriever_tool(
        retriever=retriever,
        name="company_knowledge_search",
        description=(
            "Search and retrieve relevant passages from the uploaded company documents "
            "and knowledge bases. Always use this tool to ground answers."
        ),
    )
    tools = [retriever_tool]

    # 2. LLM
    llm = _get_llm()

    # 3. Prompt for the agent
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                (
                    "You are a helpful company knowledge base assistant. "
                    "You have access to internal documents via the tool "
                    "`company_knowledge_search`. "
                    "Always call that tool before answering, and base your answer "
                    "only on retrieved content when possible.\n\n"
                    "When you respond:\n"
                    "1. Give a clear, concise answer.\n"
                    "2. At the end, add a 'Sources:' section listing each cited "
                    "document as 'Filename (Page X)'.\n"
                    "3. If the answer is not in the documents, say you don't know "
                    "rather than guessing."
                ),
            ),
            MessagesPlaceholder("chat_history"),
            ("human", "{input}"),
            MessagesPlaceholder("agent_scratchpad"),
        ]
    )

    # 4. Build the agent + executor
    agent = create_tool_calling_agent(llm=llm, tools=tools, prompt=prompt)
    executor = AgentExecutor(agent=agent, tools=tools, verbose=False)

    # 5. Attach memory
    agent_with_history = RunnableWithMessageHistory(
        executor,
        _get_session_history,
        input_messages_key="input",
        history_messages_key="chat_history",
    )
    return agent_with_history


def _build_sources_for_question(
    question: str, kb_ids: List[str], top_k: int = 5
) -> List[Dict[str, Any]]:
    """
    Simple 'side' retrieval used to build the `sources` JSON for the /ask response.
    This runs separately from the agent's tool calls but uses the same retriever.
    """
    retriever = KBVectorRetriever(vector_store=vector_store, kb_ids=kb_ids, k=top_k)
    docs = retriever.get_relevant_documents(question)

    sources: List[Dict[str, Any]] = []
    for doc in docs:
        meta = doc.metadata or {}
        sources.append(
            {
                "kb_id": meta.get("kb_id"),
                "document_id": meta.get("doc_id"),
                "filename": meta.get("filename"),
                "page": meta.get("page"),
                "score": meta.get("score"),
            }
        )
    return sources

# -----------------------------------------------------------------------------
# KB endpoints
# -----------------------------------------------------------------------------

@app.route("/kb", methods=["GET", "POST"])
def handle_kb_collection():
    """
    GET  /kb   -> list all knowledge bases
    POST /kb   -> create a new KB
    """
    if request.method == "GET":
        return jsonify({"knowledge_bases": list(knowledge_bases.values())})

    # POST - create KB
    data = request.json or {}
    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "KB 'name' is required"}), 400

    description = (data.get("description") or "").strip()
    visibility = (data.get("visibility") or "private").strip() or "private"

    kb_id = str(uuid4())
    now = _now_iso()
    knowledge_bases[kb_id] = {
        "id": kb_id,
        "name": name,
        "description": description,
        "visibility": visibility,
        "created_at": now,
        "updated_at": now,
        "document_ids": [],
    }

    return jsonify(knowledge_bases[kb_id]), 201


@app.route("/kb/<kb_id>", methods=["GET"])
def handle_single_kb(kb_id: str):
    """
    GET /kb/<kb_id> -> KB details (including document_ids)
    """
    kb = knowledge_bases.get(kb_id)
    if not kb:
        return jsonify({"error": f"Knowledge base '{kb_id}' not found"}), 404

    # Attach doc_count for convenience
    kb_with_count = dict(kb)
    kb_with_count["document_count"] = len(kb_with_count.get("document_ids", []))

    return jsonify(kb_with_count)

# -----------------------------------------------------------------------------
# Document upload & management
# -----------------------------------------------------------------------------

@app.route("/upload", methods=["POST"])
def handle_upload():
    """
    Upload one or more documents into a KB and index them into the vector DB.

    Form fields (multipart/form-data):
    - files: file(s)
    - kb_id: optional; defaults to 'default'
    - tags: optional; comma-separated string
    """
    if "files" not in request.files:
        return jsonify({"error": "No files uploaded"}), 400

    kb_id = request.form.get("kb_id", DEFAULT_KB_ID).strip() or DEFAULT_KB_ID
    if kb_id not in knowledge_bases:
        return jsonify({"error": f"Knowledge base '{kb_id}' does not exist"}), 400

    tags_raw = request.form.get("tags") or ""
    tags = [t.strip() for t in tags_raw.split(",") if t.strip()]

    uploaded_files = request.files.getlist("files")
    new_docs: List[Dict[str, Any]] = []

    for file in uploaded_files:
        if not file or file.filename == "":
            continue

        orig_filename = secure_filename(file.filename)
        file_bytes = file.read()
        if not file_bytes:
            continue

        # Basic type detection
        _, ext = os.path.splitext(orig_filename)
        ext = ext.lower()

        # Give every uploaded doc a fresh ID (no dedup for now)
        doc_id = str(uuid4())
        stored_filename = f"{doc_id}_{orig_filename}"
        stored_path = os.path.join(app.config["UPLOAD_FOLDER"], stored_filename)

        # Persist original file for later viewing
        with open(stored_path, "wb") as f:
            f.write(file_bytes)

        # Extract text pages
        try:
            if ext == ".pdf":
                pages = pdf_processor.process_pdf(stored_path)
            elif ext == ".txt":
                text = file_bytes.decode("utf-8", errors="ignore")
                pages = [
                    {
                        "text": text,
                        "metadata": {"filename": orig_filename, "page": 1},
                    }
                ]
            elif ext == ".docx":
                # Lazy import; only needed if DOCX is actually used
                from io import BytesIO
                from docx import Document as DocxDocument

                docx_obj = DocxDocument(BytesIO(file_bytes))
                text = "\n".join(p.text for p in docx_obj.paragraphs)
                pages = [
                    {
                        "text": text,
                        "metadata": {"filename": orig_filename, "page": 1},
                    }
                ]
            else:
                return (
                    jsonify(
                        {
                            "error": f"Unsupported file type '{ext}'. "
                                     "Currently supported: .pdf, .txt, .docx"
                        }
                    ),
                    400,
                )
        except Exception as e:
            app.logger.error(f"Error parsing file {orig_filename}: {e}")
            return jsonify({"error": f"Failed to process {orig_filename}"}), 500

        # Cache full pages by doc_id (for /documents/<doc_id> viewer)
        document_cache.add_document(doc_id, pages)

        # Index each page into the vector DB
        total_indexed = 0
        for page in pages:
            page_text = (page.get("text") or "").strip()
            if not page_text:
                continue

            page_meta = page.get("metadata") or {}
            page_num = page_meta.get("page", 1)

            metadata = {
                "kb_id": kb_id,
                "doc_id": doc_id,
                "filename": orig_filename,
                "page": page_num,
                "doc_text": page_text,
            }
            vector_store.add_document(page_text, metadata)
            total_indexed += 1

        now = _now_iso()
        documents[doc_id] = {
            "id": doc_id,
            "kb_id": kb_id,
            "filename": orig_filename,
            "file_type": ext.lstrip("."),
            "path": stored_path,
            "status": "ready" if total_indexed > 0 else "empty",
            "tags": tags,
            "page_count": len(pages),
            "created_at": now,
            "updated_at": now,
        }

        kb = knowledge_bases[kb_id]
        kb["document_ids"].append(doc_id)
        kb["updated_at"] = now

        new_docs.append(documents[doc_id])

    if not new_docs:
        return jsonify({"message": "No files processed", "documents": []}), 400

    return (
        jsonify(
            {
                "message": f"Uploaded {len(new_docs)} document(s) to KB '{kb_id}'",
                "documents": new_docs,
            }
        ),
        201,
    )


@app.route("/documents", methods=["GET"])
def list_documents():
    """
    GET /documents?kb_id=<optional>

    Returns a list of documents (optionally filtered by KB).
    """
    kb_id = request.args.get("kb_id")
    docs = list(documents.values())
    if kb_id:
        docs = [d for d in docs if d.get("kb_id") == kb_id]

    return jsonify({"documents": docs})


@app.route("/documents/<doc_id>", methods=["GET"])
def get_document(doc_id: str):
    """
    GET /documents/<doc_id>

    Returns:
    - document metadata
    - pages: list[{ text, metadata }]
    """
    doc = documents.get(doc_id)
    if not doc:
        return jsonify({"error": f"Document '{doc_id}' not found"}), 404

    pages = document_cache.get_document(doc_id)
    if pages is None:
        # Try to recompute from stored file as a fallback
        stored_path = doc.get("path")
        if stored_path and os.path.exists(stored_path):
            _, ext = os.path.splitext(stored_path)
            ext = ext.lower()
            try:
                if ext == ".pdf":
                    pages = pdf_processor.process_pdf(stored_path)
                elif ext == ".txt":
                    with open(stored_path, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()
                    pages = [
                        {
                            "text": text,
                            "metadata": {
                                "filename": doc.get("filename"),
                                "page": 1,
                            },
                        }
                    ]
                elif ext == ".docx":
                    from docx import Document as DocxDocument

                    docx_obj = DocxDocument(stored_path)
                    text = "\n".join(p.text for p in docx_obj.paragraphs)
                    pages = [
                        {
                            "text": text,
                            "metadata": {
                                "filename": doc.get("filename"),
                                "page": 1,
                            },
                        }
                    ]
                else:
                    pages = []
            except Exception as e:
                app.logger.error(f"Failed to reprocess {stored_path}: {e}")
                pages = []

            document_cache.add_document(doc_id, pages)
        else:
            pages = []

    return jsonify(
        {
            "document": doc,
            "pages": pages,
        }
    )

# -----------------------------------------------------------------------------
# ASK endpoint – Agentic KB Q&A
# -----------------------------------------------------------------------------

@app.route("/ask", methods=["POST"])
def handle_ask():
    """
    POST /ask

    JSON body:
    {
      "question": "string",          # required
      "kb_ids": ["kb1", "kb2"],      # optional; defaults to all KBs
      "conversation_id": "uuid",     # optional; new one created if missing
      "top_k": 5                     # optional; default 5
    }

    Response:
    {
      "answer": "string",
      "sources": [
        {
          "kb_id": "...",
          "document_id": "...",
          "filename": "...",
          "page": 1,
          "score": 0.123
        },
        ...
      ],
      "conversation_id": "uuid"
    }
    """
    data = request.json or {}
    question = (data.get("question") or "").strip()
    if not question:
        return jsonify({"error": "Field 'question' is required"}), 400

    if not documents:
        return jsonify({"error": "No documents indexed yet"}), 400

    kb_ids_in_req = data.get("kb_ids") or []
    if kb_ids_in_req:
        # Validate provided KBs
        invalid = [kb for kb in kb_ids_in_req if kb not in knowledge_bases]
        if invalid:
            return jsonify({"error": f"Unknown KB IDs: {invalid}"}), 400
        kb_ids = kb_ids_in_req
    else:
        # Default: all KBs
        kb_ids = list(knowledge_bases.keys())

    if not kb_ids:
        return jsonify({"error": "No knowledge bases available"}), 400

    try:
        top_k_raw = data.get("top_k", 5)
        top_k = max(1, int(top_k_raw))
    except (TypeError, ValueError):
        return jsonify({"error": "top_k must be an integer"}), 400

    conversation_id = data.get("conversation_id") or str(uuid4())

    # Build agent (with memory bound to conversation_id)
    agent = _build_kb_agent_with_history(kb_ids=kb_ids, top_k=top_k)

    try:
        result = agent.invoke(
            {"input": question},
            config={"configurable": {"session_id": conversation_id}},
        )
    except Exception as e:
        app.logger.error(f"/ask agent error: {e}")
        return jsonify({"error": "Agent failed to answer"}), 500

    # AgentExecutor returns a dict; "output" contains the final reply
    if isinstance(result, dict):
        answer_text = result.get("output", "")
    else:
        answer_text = str(result)

    # Build structured sources using the same KB filter
    sources = _build_sources_for_question(question, kb_ids=kb_ids, top_k=top_k)

    return jsonify(
        {
            "answer": answer_text,
            "sources": sources,
            "conversation_id": conversation_id,
        }
    )

# -----------------------------------------------------------------------------
# SEARCH endpoint – semantic search over the vector DB
# -----------------------------------------------------------------------------

@app.route("/search", methods=["POST"])
def handle_search():
    """
    POST /search

    JSON body:
    {
      "query": "string",          # required
      "kb_ids": ["kb1", "kb2"],   # optional; default all
      "top_k": 10                 # optional; default 10
    }

    Response:
    {
      "results": [
        {
          "kb_id": "...",
          "document_id": "...",
          "filename": "...",
          "page": 1,
          "score": 0.123,
          "snippet": "..."
        },
        ...
      ]
    }
    """
    data = request.json or {}
    query = (data.get("query") or "").strip()
    if not query:
        return jsonify({"error": "Field 'query' is required"}), 400

    if not documents:
        return jsonify({"error": "No documents indexed yet"}), 400

    kb_ids_in_req = data.get("kb_ids") or []
    if kb_ids_in_req:
        invalid = [kb for kb in kb_ids_in_req if kb not in knowledge_bases]
        if invalid:
            return jsonify({"error": f"Unknown KB IDs: {invalid}"}), 400
        kb_ids = kb_ids_in_req
    else:
        kb_ids = list(knowledge_bases.keys())

    try:
        top_k_raw = data.get("top_k", 10)
        top_k = max(1, int(top_k_raw))
    except (TypeError, ValueError):
        return jsonify({"error": "top_k must be an integer"}), 400

    retriever = KBVectorRetriever(vector_store=vector_store, kb_ids=kb_ids, k=top_k)
    docs = retriever.get_relevant_documents(query)

    results: List[Dict[str, Any]] = []
    for doc in docs:
        meta = doc.metadata or {}
        results.append(
            {
                "kb_id": meta.get("kb_id"),
                "document_id": meta.get("doc_id"),
                "filename": meta.get("filename"),
                "page": meta.get("page"),
                "score": meta.get("score"),
                "snippet": doc.page_content,
            }
        )

    return jsonify({"results": results})

# -----------------------------------------------------------------------------
# RESET endpoint – wipe in-memory state
# -----------------------------------------------------------------------------

@app.route("/reset", methods=["POST"])
def handle_reset():
    """
    Reset in-memory state:
    - vector_store
    - knowledge_bases (recreates default)
    - documents
    - document_cache
    - session histories
    """
    global vector_store, document_cache, _session_histories

    vector_store = VectorStore()
    document_cache = DocumentCache(ttl=3600)
    documents.clear()
    knowledge_bases.clear()
    _session_histories.clear()
    _ensure_default_kb()

    return jsonify({"message": "System reset successfully"})


# -----------------------------------------------------------------------------
# Entrypoint
# -----------------------------------------------------------------------------

if __name__ == "__main__":
    # Ensure upload folder exists
    os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
    app.run(host="0.0.0.0", port=5000, debug=True)
